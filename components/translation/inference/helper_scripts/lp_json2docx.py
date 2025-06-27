import json
import os
from docx import Document
from pprint import pformat

def json_to_word(json_path, doc_path, resources):
    
    # Load JSON data from the file
    with open(json_path, 'r') as file:
        data = json.load(file)
    
    new_data = {}
    for key, value in data.items():
        if key in resources: new_data[key] = value
    data = new_data

    # Create a new Word Document
    doc = Document()
    doc = add_content(doc, data, level=1)

    doc.save(doc_path)

def add_content(doc, data, level):
    if isinstance(data, dict):
        for key, value in data.items():
            doc.add_heading(key, level=level)
            add_content(doc, value, level+1)
    elif isinstance(data, list):
        for item in data:
            add_content(doc, item, level)
    elif isinstance(data, str):
        doc.add_paragraph(data.replace("\\n", "\n"))

    return doc


if __name__ == "__main__":
    # Define the path to the JSON file and the output Word document
    filename = "temp"
    json_path = f'./{filename}.json'
    json_path = "/home/t-narora/translation/eng_lps/sample_lp.json"
    doc_path = f'./{filename}_og.docx'
    resources = [
        "_id",
        "extracted_resources",
        "additional_resources",
        "checklist"
    ]
    # Convert JSON to Word Document
    json_to_word(json_path, doc_path, resources)

    print(f"Document saved to {doc_path}") 